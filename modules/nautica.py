"""
nautica.py — Генератор отчётов Word  v4
SSK Zvezda | The First Whistle
Деловой формат ССК Звезда — выравнивание по ширине, детальная информация
"""

import sys, logging
from datetime import datetime
from pathlib import Path

log = logging.getLogger("nautica")

def get_data_path():
    if getattr(sys,"frozen",False): return Path(sys.executable).parent
    return Path(__file__).parent.parent

DATA_PATH = get_data_path()

# ── Цвета (RGB) для Word ───────────────────────────────────────────────────────
CLR_BLACK  = (0x1A, 0x1A, 0x1A)
CLR_DKGRAY = (0x44, 0x44, 0x44)
CLR_GRAY   = (0x77, 0x77, 0x77)
CLR_LGRAY  = (0xBB, 0xBB, 0xBB)
CLR_WHITE  = (0xFF, 0xFF, 0xFF)
CLR_OK     = (0x1A, 0x7A, 0x3A)
CLR_WARN   = (0x8A, 0x5A, 0x00)
CLR_CRIT   = (0xAA, 0x1A, 0x1A)

def _rgb(r,g,b):
    from docx.shared import RGBColor
    return RGBColor(r,g,b)

def generate_report(period, dept_avg, kpi_summary, employees, resources,
                    low_resources, output_path=None, output_dir=None,
                    template="official", kpi_red_zone=40):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import matplotlib; matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from io import BytesIO

    # ── Путь сохранения ────────────────────────────────────────────────────────
    if output_path is None:
        ts  = datetime.now().strftime("%Y%m%d_%H%M%S")
        fn  = f"SSK_Zvezda_KPI_{period}_{ts}.docx"
        # Используем output_dir если передан, иначе DATA_PATH
        save_dir = Path(output_dir) if output_dir else DATA_PATH
        # Проверяем что директория существует
        if not save_dir.exists():
            try:
                save_dir.mkdir(parents=True, exist_ok=True)
            except Exception:
                save_dir = DATA_PATH
        output_path = str(save_dir / fn)

    doc = Document()

    # ── Поля страницы ──────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(1.5)
        sec.page_width    = Cm(21.0)
        sec.page_height   = Cm(29.7)

    # ── Стили параграфов ───────────────────────────────────────────────────────
    def _para(text="", bold=False, italic=False, size=11, color=CLR_BLACK,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY, font="Arial"):
        p   = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.name     = font
        run.font.size     = Pt(size)
        run.font.bold     = bold
        run.font.italic   = italic
        run.font.color.rgb= _rgb(*color)
        pf = p.paragraph_format
        pf.space_after    = Pt(4)
        pf.space_before   = Pt(2)
        return p, run

    def _heading(text, lvl=1):
        sizes = {1:16, 2:13, 3:11}
        p, run = _para(text, bold=True, size=sizes.get(lvl,12),
                       color=CLR_BLACK,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(10 if lvl==1 else 6)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def _divider():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),"single")
        bottom.set(qn("w:sz"),"6")
        bottom.set(qn("w:space"),"1")
        bottom.set(qn("w:color"),"1A1A1A")
        pBdr.append(bottom); pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(4)

    def _cell_shade(cell, hex_color="F5F5F5"):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _cell_text(cell, text, bold=False, size=10, color=CLR_BLACK,
                   align=WD_ALIGN_PARAGRAPH.LEFT, font="Arial"):
        cell.text = ""
        p   = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.name     = font
        run.font.size     = Pt(size)
        run.font.bold     = bold
        run.font.color.rgb= _rgb(*color)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)

    def _set_col_widths(table, widths_cm):
        from docx.shared import Cm
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths_cm):
                    cell.width = Cm(widths_cm[i])

    # ══════════════════════════════════════════════════════════════════════════
    # ШАПКА ДОКУМЕНТА
    # ══════════════════════════════════════════════════════════════════════════
    # Логотипная строка
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_p.add_run("ООО «ССК ЗВЕЗДА»")
    run.font.name = "Arial"; run.font.size = Pt(14); run.font.bold = True
    run.font.color.rgb = _rgb(*CLR_BLACK)
    logo_p.paragraph_format.space_after = Pt(0)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Отдел кадров и управление производительностью")
    run.font.name = "Arial"; run.font.size = Pt(10)
    run.font.color.rgb = _rgb(*CLR_DKGRAY)
    sub_p.paragraph_format.space_after = Pt(2)

    _divider()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"АНАЛИТИЧЕСКИЙ ОТЧЁТ ПО KPI\nЗА ПЕРИОД: {period}")
    run.font.name = "Arial"; run.font.size = Pt(15); run.font.bold = True
    run.font.color.rgb = _rgb(*CLR_BLACK)
    title_p.paragraph_format.space_after = Pt(6)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}  "
                          f"  Время: {datetime.now().strftime('%H:%M')}"
                          f"  |  Версия: 4.0")
    run.font.name = "Arial"; run.font.size = Pt(9)
    run.font.color.rgb = _rgb(*CLR_GRAY)
    _divider()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. ИСПОЛНИТЕЛЬНОЕ РЕЗЮМЕ
    # ══════════════════════════════════════════════════════════════════════════
    _heading("1. ИСПОЛНИТЕЛЬНОЕ РЕЗЮМЕ", 1)

    avg_kpi = round(sum(dept_avg.values())/len(dept_avg),1) if dept_avg else 0
    active  = sum(1 for e in employees if e.get("active",1))
    inactive= len(employees) - active
    crit_kpi= sum(1 for k in kpi_summary if k.get("avg_score",100) < kpi_red_zone)

    # Сводная таблица показателей
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(lbl,val) in enumerate([
        ("Всего сотрудников",    str(len(employees))),
        ("Активных / Архив",     f"{active} / {inactive}"),
        ("Средний KPI",          f"{avg_kpi:.1f}"),
        ("Критических KPI",      str(crit_kpi)),
    ]):
        _cell_shade(tbl.rows[0].cells[i], "1A1A1A")
        _cell_text(tbl.rows[0].cells[i], lbl, bold=True, size=10,
                   color=CLR_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    row2 = tbl.add_row()
    for i,(lbl,val) in enumerate([
        ("Всего сотрудников",    str(len(employees))),
        ("Активных / Архив",     f"{active} / {inactive}"),
        ("Средний KPI",          f"{avg_kpi:.1f}"),
        ("Критических KPI",      str(crit_kpi)),
    ]):
        _cell_shade(row2.cells[i], "F5F5F5")
        clr = CLR_CRIT if (lbl=="Средний KPI" and avg_kpi<kpi_red_zone) else CLR_BLACK
        _cell_text(row2.cells[i], val, bold=True, size=14,
                   color=clr, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_col_widths(tbl, [3.75, 3.75, 3.75, 3.75])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Аналитический вывод
    level = "высокий" if avg_kpi>=80 else ("критически низкий" if avg_kpi<kpi_red_zone else "средний")
    _para(f"По итогам анализа периода {period} общий уровень производительности предприятия "
          f"оценивается как {level} (средний KPI: {avg_kpi:.1f} балла). "
          f"Из {len(employees)} сотрудников предприятия {active} являются активными. "
          f"Выявлено {crit_kpi} записей с показателями ниже установленной красной зоны ({kpi_red_zone} баллов).")

    # Ресурсы
    _para(f"Состояние складских запасов: из {len(resources)} позиций {len(low_resources)} "
          f"находятся на уровне или ниже минимального порога, что требует оперативного "
          f"пополнения.")

    # ══════════════════════════════════════════════════════════════════════════
    # 2. ГРАФИК KPI ПО ОТДЕЛАМ
    # ══════════════════════════════════════════════════════════════════════════
    if dept_avg:
        _heading("2. ПОКАЗАТЕЛИ KPI ПО ПОДРАЗДЕЛЕНИЯМ", 1)
        items  = sorted(dept_avg.items(), key=lambda x:x[1], reverse=True)
        depts  = [x[0] for x in items]; scores = [x[1] for x in items]
        colors = [f"#{CLR_OK[0]:02x}{CLR_OK[1]:02x}{CLR_OK[2]:02x}" if s>=80
                  else (f"#{CLR_CRIT[0]:02x}{CLR_CRIT[1]:02x}{CLR_CRIT[2]:02x}" if s<kpi_red_zone
                        else "#555555") for s in scores]

        fig,ax = plt.subplots(figsize=(14,3), facecolor="white")
        ax.set_facecolor("white"); ax.set_axisbelow(True)
        ax.xaxis.grid(True,color="#EEEEEE",linewidth=0.8)
        bars = ax.barh(depts, scores, color=colors, edgecolor="white", height=0.52)
        for bar,s in zip(bars,scores):
            ax.text(bar.get_width()+0.4, bar.get_y()+bar.get_height()/2,
                    f"{s:.1f}", va="center", ha="left",
                    color="#444444", fontsize=10, fontweight="bold", fontfamily="Arial")
        ax.set_xlim(0,110); ax.axvline(80,color="#AAAAAA",linestyle="--",linewidth=1.5,alpha=0.9)
        ax.text(80.5, len(depts)-0.05,"норма 80",color="#AAAAAA",fontsize=9,va="top",fontfamily="Arial")
        if kpi_red_zone:
            ax.axvline(kpi_red_zone,color="#DDAAAA",linestyle=":",linewidth=1)
            ax.text(kpi_red_zone+0.4,0.05,f"зона риска",color="#DDAAAA",fontsize=8,va="bottom")
        ax.set_xlabel("KPI, баллов",color="#555555",fontsize=10,fontfamily="Arial")
        ax.tick_params(colors="#555555",labelsize=11); ax.tick_params(axis='y',labelsize=12)
        for sp in ["top","right","left"]: ax.spines[sp].set_visible(False)
        ax.spines["bottom"].set_color("#CCCCCC"); plt.tight_layout(pad=0.5)
        buf=BytesIO(); fig.savefig(buf,format="png",dpi=140,bbox_inches="tight",facecolor="white")
        plt.close(fig); buf.seek(0)
        doc.add_picture(buf, width=Cm(15))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().paragraph_format.space_after=Pt(4)

        # Таблица детализации
        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.style = "Table Grid"; tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci,(h) in enumerate(["Подразделение","Средний KPI","Статус"]):
            _cell_shade(tbl2.rows[0].cells[ci],"1A1A1A")
            _cell_text(tbl2.rows[0].cells[ci],h,bold=True,size=10,
                       color=CLR_WHITE,align=WD_ALIGN_PARAGRAPH.CENTER)
        for d,s in items:
            r3=tbl2.add_row()
            status = "НОРМА" if s>=80 else ("РИСК" if s<kpi_red_zone else "СРЕДНИЙ")
            clr    = CLR_OK if s>=80 else (CLR_CRIT if s<kpi_red_zone else CLR_WARN)
            shade  = "F5F5F5" if list(items).index((d,s))%2==0 else "FFFFFF"
            _cell_shade(r3.cells[0],shade); _cell_shade(r3.cells[1],shade); _cell_shade(r3.cells[2],shade)
            _cell_text(r3.cells[0],d,size=10,color=CLR_BLACK)
            _cell_text(r3.cells[1],f"{s:.1f}",size=11,bold=True,
                       color=clr,align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(r3.cells[2],status,size=10,bold=True,
                       color=clr,align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_col_widths(tbl2,[8,3.5,3.5])
        doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # 3. РЕЙТИНГ СОТРУДНИКОВ — ТОП-20
    # ══════════════════════════════════════════════════════════════════════════
    _heading("3. РЕЙТИНГ СОТРУДНИКОВ ПО KPI (ТОП-20)", 1)
    top20 = sorted(kpi_summary, key=lambda x:x.get("avg_score",0), reverse=True)[:20]

    tbl3 = doc.add_table(rows=1, cols=5)
    tbl3.style = "Table Grid"; tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci,h in enumerate(["№","ФИО","Подразделение","KPI","Задач"]):
        _cell_shade(tbl3.rows[0].cells[ci],"1A1A1A")
        _cell_text(tbl3.rows[0].cells[ci],h,bold=True,size=10,
                   color=CLR_WHITE,align=WD_ALIGN_PARAGRAPH.CENTER)
    for i,r in enumerate(top20):
        rrow = tbl3.add_row()
        sc   = r.get("avg_score",0)
        clr  = CLR_OK if sc>=80 else (CLR_CRIT if sc<kpi_red_zone else CLR_WARN)
        shade= "F0F0F0" if i%2==0 else "FFFFFF"
        for ci in range(5): _cell_shade(rrow.cells[ci],shade)
        _cell_text(rrow.cells[0],str(i+1),size=9,color=CLR_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(rrow.cells[1],r.get("full_name",""),size=10,color=CLR_BLACK)
        _cell_text(rrow.cells[2],r.get("department",""),size=9,color=CLR_DKGRAY)
        _cell_text(rrow.cells[3],f"{sc:.1f}",size=11,bold=True,
                   color=clr,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(rrow.cells[4],str(r.get("total_tasks",0)),size=10,
                   color=CLR_DKGRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_col_widths(tbl3,[1.0,5.5,3.5,2.0,1.5])
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # 4. КРИТИЧЕСКИЕ ПОКАЗАТЕЛИ
    # ══════════════════════════════════════════════════════════════════════════
    crit_list = [k for k in kpi_summary if k.get("avg_score",100) < kpi_red_zone]
    if crit_list:
        _heading(f"4. СОТРУДНИКИ В ЗОНЕ РИСКА (KPI < {kpi_red_zone})", 1)
        _para(f"Выявлено {len(crit_list)} сотрудников с показателями ниже порогового значения. "
              f"Рекомендуется провести индивидуальные беседы и разработать план корректирующих мероприятий.",
              color=CLR_CRIT)

        tbl4=doc.add_table(rows=1,cols=4)
        tbl4.style="Table Grid"; tbl4.alignment=WD_TABLE_ALIGNMENT.CENTER
        for ci,h in enumerate(["ФИО","Подразделение","Период","KPI"]):
            _cell_shade(tbl4.rows[0].cells[ci],"AA1A1A")
            _cell_text(tbl4.rows[0].cells[ci],h,bold=True,size=10,
                       color=CLR_WHITE,align=WD_ALIGN_PARAGRAPH.CENTER)
        for i,r in enumerate(sorted(crit_list,key=lambda x:x.get("avg_score",0))):
            rrow=tbl4.add_row()
            shade="FFF0F0" if i%2==0 else "FFFFFF"
            for ci in range(4): _cell_shade(rrow.cells[ci],shade)
            _cell_text(rrow.cells[0],r.get("full_name",""),size=10,color=CLR_BLACK)
            _cell_text(rrow.cells[1],r.get("department",""),size=9,color=CLR_DKGRAY)
            _cell_text(rrow.cells[2],r.get("period",""),size=10,color=CLR_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(rrow.cells[3],f"{r.get('avg_score',0):.1f}",size=11,bold=True,
                       color=CLR_CRIT,align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_col_widths(tbl4,[5.5,3.5,2.5,2.0])
        doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # 5. СКЛАДСКИЕ ЗАПАСЫ
    # ══════════════════════════════════════════════════════════════════════════
    _heading("5. СОСТОЯНИЕ СКЛАДСКИХ ЗАПАСОВ", 1)

    if low_resources:
        _para(f"⚠  ВНИМАНИЕ: {len(low_resources)} позиций склада находятся на уровне или ниже "
              f"минимального порога и требуют незамедлительного пополнения.",
              bold=True, color=CLR_CRIT)
        doc.add_paragraph().paragraph_format.space_after=Pt(2)

    tbl5=doc.add_table(rows=1,cols=6)
    tbl5.style="Table Grid"; tbl5.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ci,h in enumerate(["Наименование","Категория","Ед.","Остаток","Минимум","Статус"]):
        _cell_shade(tbl5.rows[0].cells[ci],"1A1A1A")
        _cell_text(tbl5.rows[0].cells[ci],h,bold=True,size=10,
                   color=CLR_WHITE,align=WD_ALIGN_PARAGRAPH.CENTER)

    for i,r in enumerate(resources):
        rrow=tbl5.add_row()
        is_low=r["quantity"]<=r["min_quantity"]
        shade="FFF0F0" if is_low else ("F5F5F5" if i%2==0 else "FFFFFF")
        for ci in range(6): _cell_shade(rrow.cells[ci],shade)
        _cell_text(rrow.cells[0],r["name"],size=10,color=CLR_BLACK)
        _cell_text(rrow.cells[1],r["category"],size=9,color=CLR_DKGRAY)
        _cell_text(rrow.cells[2],r["unit"],size=10,color=CLR_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
        q_clr=CLR_CRIT if is_low else CLR_OK
        _cell_text(rrow.cells[3],str(r["quantity"]),size=10,bold=is_low,
                   color=q_clr,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(rrow.cells[4],str(r["min_quantity"]),size=10,
                   color=CLR_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
        status="⚠ НИЖЕ НОРМЫ" if is_low else "✓ Норма"
        _cell_text(rrow.cells[5],status,size=9,bold=is_low,
                   color=q_clr,align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_col_widths(tbl5,[4.5,2.5,1.0,1.5,1.5,2.0])
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # 6. КАДРОВЫЙ СОСТАВ
    # ══════════════════════════════════════════════════════════════════════════
    _heading("6. КАДРОВЫЙ СОСТАВ ПРЕДПРИЯТИЯ", 1)

    from collections import Counter, defaultdict
    by_dept  = Counter(e["department"] for e in employees)
    by_pos   = defaultdict(list)
    for e in employees: by_pos[e["department"]].append(e["position"])

    tbl6=doc.add_table(rows=1,cols=4)
    tbl6.style="Table Grid"; tbl6.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ci,h in enumerate(["Подразделение","Сотрудников","% от общего","Ср. стаж"]):
        _cell_shade(tbl6.rows[0].cells[ci],"1A1A1A")
        _cell_text(tbl6.rows[0].cells[ci],h,bold=True,size=10,
                   color=CLR_WHITE,align=WD_ALIGN_PARAGRAPH.CENTER)
    total=len(employees) or 1
    for i,(dept,cnt) in enumerate(sorted(by_dept.items())):
        rrow=tbl6.add_row()
        shade="F5F5F5" if i%2==0 else "FFFFFF"
        for ci in range(4): _cell_shade(rrow.cells[ci],shade)
        # Посчитаем средний стаж в годах
        dept_emps=[e for e in employees if e["department"]==dept]
        stazh_years=[]
        for e in dept_emps:
            try:
                hd=datetime.strptime(str(e["hire_date"])[:10],"%Y-%m-%d")
                stazh_years.append((datetime.now()-hd).days/365)
            except: pass
        avg_stazh=f"{sum(stazh_years)/len(stazh_years):.1f} лет" if stazh_years else "—"
        _cell_text(rrow.cells[0],dept,size=10,color=CLR_BLACK)
        _cell_text(rrow.cells[1],str(cnt),size=11,bold=True,
                   color=CLR_BLACK,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(rrow.cells[2],f"{cnt/total*100:.1f}%",size=10,
                   color=CLR_DKGRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(rrow.cells[3],avg_stazh,size=10,
                   color=CLR_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_col_widths(tbl6,[6.0,2.5,2.5,2.5])
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # ПОДПИСИ
    # ══════════════════════════════════════════════════════════════════════════
    _divider()
    sign_tbl=doc.add_table(rows=2,cols=2)
    for row in sign_tbl.rows:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr()
    sign_tbl.rows[0].cells[0].text=""
    sign_tbl.rows[0].cells[1].text=""
    _cell_text(sign_tbl.rows[1].cells[0],
               "Подготовил: _______________  /_______________/",
               size=10,color=CLR_DKGRAY)
    _cell_text(sign_tbl.rows[1].cells[1],
               "Утвердил: _______________  /_______________/",
               size=10,color=CLR_DKGRAY,align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(output_path)
    log.info("Отчёт сохранён: %s", output_path)
    return output_path
